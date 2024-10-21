import os
import pandas as pd
from docx import Document
from docx2pdf import convert
from datetime import datetime

# Load the CSV file with the mapping of old names (FirstName_LastName) to new names (Name)
csv_file = 'completed_modules.csv'
df = pd.read_csv(csv_file)

# Load the Word document
docx_file = 'certificate_template.docx'

# Specify the path to the Downloads folder
downloads_folder = "Certificates"

# Create the 'Certificates' folder if it doesn't exist
if not os.path.exists(downloads_folder):
    os.makedirs(downloads_folder)

# Prompt user for input
challenge_name = input("Enter the name of the challenge: ")
mlsa_name = input("Enter the name of MLSA: ")
rank = input("Enter the rank of MLSA (alpha/beta/gold): ")

# Prompt MLSA for the date in dd month year format
date = input("Enter the date in dd month year format: ")

# Iterate through each student's name in the CSV file
for index, row in df.iterrows():
    student_name = row['Name']

    # Make a copy of the certificate template for each student
    doc = Document(docx_file)

    # Replace the placeholder text with user-provided values
    placeholders = {
        'FirstName_LastName': student_name,
        'ChallengeName': challenge_name,
        'Date': date,
        'MLSAName': mlsa_name,
        'Rank': rank
    }

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for placeholder, value in placeholders.items():
                run.text = run.text.replace(placeholder, value)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for placeholder, value in placeholders.items():
                            run.text = run.text.replace(placeholder, value)

    # Save the edited certificate as a temporary .docx file in the Downloads folder
    temp_docx_filename = os.path.join(downloads_folder, f'{student_name}_Certificate.docx')
    doc.save(temp_docx_filename)

    # Convert the temporary .docx file to PDF using docx2pdf
    pdf_filename = os.path.join(downloads_folder, f'{student_name}_Certificate.pdf')
    convert(temp_docx_filename, pdf_filename)

print("Certificates have been generated in PDF format for all students in the same folder.")
